import random
import datetime
import os
from pprint import pprint

from pandas_ods_reader import read_ods
import pandas as pd
import docx
import json
import numpy as np

import params as p
import analysis as a

def shorten_sample_path(inpath):
    path = inpath
    newpath = []
    el = None
    while el != '' and os.path.normpath(path) != os.path.normpath(p.p_input):
        path, el = os.path.split(path)
        newpath.insert(0, el)
    return os.path.join(*newpath[1:])

def print_report(report, filename):
    jsonfile = filename + ".json"
    json.dump(report, open(jsonfile, "wt"), indent=4, ensure_ascii=False )

    doc = docx.Document()

    doc.add_heading('Визуализция выборок', 0)
    for i, s in enumerate(report['normality']): # выборка
        if i > 0:
            doc.add_paragraph(' ')
            doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)

        doc.add_heading(f'Выборка {s["sample"]}:', 1)
        for subs in s['data']: # Подвыборки
            for obj in subs['results']: # объекты
                doc.add_heading(f'Подвыборка {shorten_sample_path(subs["path"])}, объект {obj["object"]}:', 2)
                doc.add_picture(obj["grafic"], width=docx.shared.Cm(12.5))

    # ========
    # Результаты тестов Шапиро (TRUE / FALSE)
    doc.add_paragraph(' ')
    doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    doc.add_heading('Тесты на нормальность', 0)
    doc.add_heading('Тест Шапиро-Уилка (РЕШЕНИЕ)', 1)
    for s in report['normality']:  # выборка
        doc.add_heading(f'Выборка {s["sample"]}:', 2)
        r = len(s['data']) + 1  # Number of rows you want
        c = len(p.n_objects_to_analysis) + 1  # Number of collumns you want
        table = doc.add_table(rows=r, cols=c, style='LightShading-Accent1')

        for y in range(1, c):
            cell = table.cell(0, y)
            cell.text = p.n_objects_to_analysis[y - 1]

        for x, subs in enumerate(s['data']):
            path = shorten_sample_path(subs['path'])

            cell = table.cell(x + 1, 0)
            cell.text = path
            for obj in subs['results']:  # объекты
                y = p.n_objects_to_analysis.index(obj['object'])
                cell = table.cell(x + 1, y + 1)
                cell.text = 'True' if obj['normal']['shapiro']['status'] else 'False'

    # ========
    # p-value
    doc.add_heading('Тест Шапиро-Уилка (p-value)', 1)
    for s in report['normality']:  # выборка
        doc.add_heading(f'Выборка {s["sample"]}:', 2)
        r = len(s['data']) + 1  # Number of rows you want
        c = len(p.n_objects_to_analysis) + 1  # Number of collumns you want
        table = doc.add_table(rows=r, cols=c, style='LightShading-Accent1')

        for y in range(1, c):
            cell = table.cell(0, y)
            cell.text = p.n_objects_to_analysis[y - 1]

        for x, subs in enumerate(s['data']):
            path = shorten_sample_path(subs['path'])

            cell = table.cell(x + 1, 0)
            cell.text = path
            for obj in subs['results']:  # объекты
                y = p.n_objects_to_analysis.index(obj['object'])
                cell = table.cell(x + 1, y + 1)
                cell.text = str(round(obj['normal']['shapiro']['pvalue'], p.i_rounddig))

    # ========
    # Результаты тестов Андерсона-Дарлинга (TRUE / FALSE)
    doc.add_paragraph(' ')
    doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    doc.add_heading('Тест Андерсона-Дарлинга (РЕШЕНИЕ)', 1)
    for s in report['normality']:  # выборка
        doc.add_heading(f'Выборка {s["sample"]}:', 2)
        r = len(s['data']) + 1  # Number of rows you want
        c = len(p.n_objects_to_analysis) + 1  # Number of collumns you want
        table = doc.add_table(rows=r, cols=c, style='LightShading-Accent1')

        for y in range(1, c):
            cell = table.cell(0, y)
            cell.text = p.n_objects_to_analysis[y - 1]

        for x, subs in enumerate(s['data']):
            path = shorten_sample_path(subs['path'])

            cell = table.cell(x + 1, 0)
            cell.text = path
            for obj in subs['results']:  # объекты
                y = p.n_objects_to_analysis.index(obj['object'])
                cell = table.cell(x + 1, y + 1)
                cell.text = 'True' if obj['normal']['anderson']['status'] else 'False'

    # ========
    # p-value
    doc.add_heading('Тест Андерсона-Дарлинга (p-value)', 1)
    for s in report['normality']:  # выборка
        doc.add_heading(f'Выборка {s["sample"]}:', 2)
        r = len(s['data']) + 1  # Number of rows you want
        c = len(p.n_objects_to_analysis) + 1  # Number of collumns you want
        table = doc.add_table(rows=r, cols=c, style='LightShading-Accent1')

        for y in range(1, c):
            cell = table.cell(0, y)
            cell.text = p.n_objects_to_analysis[y - 1]

        for x, subs in enumerate(s['data']):
            path = shorten_sample_path(subs['path'])

            cell = table.cell(x + 1, 0)
            cell.text = path
            for obj in subs['results']:  # объекты
                y = p.n_objects_to_analysis.index(obj['object'])
                cell = table.cell(x + 1, y + 1)
                cell.text = str(round(obj['normal']['anderson']['pvalue'], p.i_rounddig))

    # ========
    # сравнение по средним
    doc.add_paragraph(' ')
    doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    doc.add_heading('Визуализция сравнений данных', 0)
    for i, s in enumerate(report['comparisons']): # выборка
        if i > 0:
            doc.add_paragraph(' ')
            doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)

        doc.add_heading(f'Выборки "{s["sample1"]}" и "{s["sample2"]}":', 1)
        for subs in s['data']: # Подвыборки
            if len(subs['results']) == 0:
                continue
            doc.add_heading(f'Подвыборки {shorten_sample_path(subs["path1"])}:', 2)
            for obj in subs['results']: # объекты
                doc.add_heading(f'Объект "{obj["object"]}":', 3)
                doc.add_picture(obj["grafic"], width=docx.shared.Cm(12.5))

    # ========
    # Результаты тестов Стьюдента (TRUE / FALSE)
    doc.add_paragraph(' ')
    doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    doc.add_heading('Тесты сравнения средних', 0)
    doc.add_heading('Сравненее средних по Стьюденту (РЕШЕНИЕ)', 1)
    for s in report['comparisons']:  # выборка
        doc.add_heading(f'Выборки "{s["sample1"]}" и "{s["sample2"]}":', 2)
        r = len(s['data']) + 1  # Number of rows you want
        c = len(p.n_objects_to_analysis) + 1  # Number of collumns you want
        table = doc.add_table(rows=r, cols=c, style='LightShading-Accent1')

        for y in range(1, c):
            cell = table.cell(0, y)
            cell.text = p.n_objects_to_analysis[y - 1]

        for x, subs in enumerate(s['data']):
            path = shorten_sample_path(subs['path1'])

            cell = table.cell(x + 1, 0)
            cell.text = path
            for obj in subs['results']:  # объекты
                if 'meantest' not in obj.keys(): continue
                y = p.n_objects_to_analysis.index(obj['object'])
                cell = table.cell(x + 1, y + 1)
                cell.text = 'True' if obj['meantest']['student']['status'] else 'False'

    # ========
    # p-value
    doc.add_heading('Сравненее средних по Стьюденту (p-value)', 1)
    for s in report['comparisons']:  # выборка
        doc.add_heading(f'Выборки "{s["sample1"]}" и "{s["sample2"]}":', 2)
        r = len(s['data']) + 1  # Number of rows you want
        c = len(p.n_objects_to_analysis) + 1  # Number of collumns you want
        table = doc.add_table(rows=r, cols=c, style='LightShading-Accent1')

        for y in range(1, c):
            cell = table.cell(0, y)
            cell.text = p.n_objects_to_analysis[y - 1]

        for x, subs in enumerate(s['data']):
            path = shorten_sample_path(subs['path1'])

            cell = table.cell(x + 1, 0)
            cell.text = path
            for obj in subs['results']:  # объекты
                if 'meantest' not in obj.keys(): continue
                y = p.n_objects_to_analysis.index(obj['object'])
                cell = table.cell(x + 1, y + 1)
                cell.text = str(round(obj['meantest']['student']['pvalue'], p.i_rounddig))

    # ========
    # Результаты тестов Манна-Уитни (TRUE / FALSE)
    doc.add_paragraph(' ')
    doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    doc.add_heading('Сравненее средних по Манну-Уитни (РЕШЕНИЕ)', 1)
    for s in report['comparisons']:  # выборка
        doc.add_heading(f'Выборки "{s["sample1"]}" и "{s["sample2"]}":', 2)
        r = len(s['data']) + 1  # Number of rows you want
        c = len(p.n_objects_to_analysis) + 1  # Number of collumns you want
        table = doc.add_table(rows=r, cols=c, style='LightShading-Accent1')

        for y in range(1, c):
            cell = table.cell(0, y)
            cell.text = p.n_objects_to_analysis[y - 1]

        for x, subs in enumerate(s['data']):
            path = shorten_sample_path(subs['path1'])

            cell = table.cell(x + 1, 0)
            cell.text = path
            for obj in subs['results']:  # объекты
                if 'meantest' not in obj.keys(): continue
                y = p.n_objects_to_analysis.index(obj['object'])
                cell = table.cell(x + 1, y + 1)
                cell.text = 'True' if obj['meantest']['mannweatney']['status'] else 'False'

    # ========
    # p-value
    doc.add_heading('Сравненее средних по Манну-Уитни (p-value)', 1)
    for s in report['comparisons']:  # выборка
        doc.add_heading(f'Выборки "{s["sample1"]}" и "{s["sample2"]}":', 2)
        r = len(s['data']) + 1  # Number of rows you want
        c = len(p.n_objects_to_analysis) + 1  # Number of collumns you want
        table = doc.add_table(rows=r, cols=c, style='LightShading-Accent1')

        for y in range(1, c):
            cell = table.cell(0, y)
            cell.text = p.n_objects_to_analysis[y - 1]

        for x, subs in enumerate(s['data']):
            path = shorten_sample_path(subs['path1'])

            cell = table.cell(x + 1, 0)
            cell.text = path
            for obj in subs['results']:  # объекты
                if 'meantest' not in obj.keys(): continue
                y = p.n_objects_to_analysis.index(obj['object'])
                cell = table.cell(x + 1, y + 1)
                cell.text = str(round(obj['meantest']['mannweatney']['pvalue'], p.i_rounddig))

    # ========
    doc.save(filename)  # Save document


def get_new_filename(dir=None):
    filename = datetime.datetime.today().\
                   strftime("%Y-%m-%d.%H.%M.%S") + \
                    f".{random.randint(0, 99999)}.png"

    if dir is not None:
        filename = os.path.join(dir, filename)

    return filename

def load_dataframe_from_folder(dir):
    data = []

    for in_file in os.listdir(dir):
        in_filepath = os.path.join(dir, in_file)
        if os.path.isdir(in_filepath):
            continue

        filename_wo_ext, ext = os.path.splitext(in_file)
        if ext not in ['.ods']:
            continue

        #print(f"Обработка файла {in_filepath}")

        try:
            df_orig = read_ods(in_filepath, p.sheet_name)
            df = df_orig[df_orig[p.n_objtype].isin(p.n_objects_to_analysis)]
            
            # проверка данных на логическую структуру
            currow = None
            max_nucleori = 0
            flag_first_nucleos = True
            prew_cytoplasm = None
            error_idx = {
                'no_nucleos': [],
                'no_nucleori': [],
                'additional_nucleos': [],
                'nan': []
            }

            for index, row in df.iterrows():
                if row[p.n_objtype] == p.n_cytoplasm:
                    if currow is not None:
                        if currow[1] is None:
                            p.logger.warning(f"[{in_filepath}] Цитоплазма без ядер! Строка [{prew_cytoplasm+2}]")
                            error_idx['no_nucleos'].append(prew_cytoplasm)
                        elif len(currow[2]) == 0:
                            p.logger.warning(f"[{in_filepath}] Клетка без ядрышек! Строка [{prew_cytoplasm+2}]")
                            error_idx['no_nucleori'].append(prew_cytoplasm)

                    currow = [None, None, []]
                    try:
                        currow[0] = float(row[p.n_value])
                    except ValueError:
                        p.logger.warning(f"[{in_filepath}] Нечисловое значение в цитоплазме! Строка [{index + 2}]")
                        error_idx['nan'].append(index+2)

                    curr_nucleori = 0
                    flag_first_nucleos = True
                    prew_cytoplasm = index

                elif row[p.n_objtype] == p.n_nucleos:
                    if flag_first_nucleos:
                        try:
                            currow[1] = float(row[p.n_value])
                        except ValueError:
                            p.logger.warning(f"[{in_filepath}] Нечисловое значение ядра! Строка [{index + 2}]")
                            error_idx['nan'].append(index)
                        flag_first_nucleos = False
                    else:
                        p.logger.warning(f"[{in_filepath}] Второе ядро в клетке! Строка [{index+2}]")
                        error_idx['additional_nucleos'].append(index)

                elif row[p.n_objtype] == p.n_nucleori:
                    try:
                        currow[2].append(float(row[p.n_value]))
                    except ValueError:
                        p.logger.warning(f"[{in_filepath}] Нечисловое значение ядрышка! Строка [{index + 2}]")
                        error_idx['nan'].append(index)
                    curr_nucleori += 1
                    if max_nucleori < curr_nucleori:
                        max_nucleori = curr_nucleori

            # # ====
            # # проверка данных на выбросы
            #
            # df_orig['chauvenet'] = ''
            # for obj in p.n_objects_to_analysis:
            #     df_err = df_orig[df_orig[p.n_objtype] == obj]
            #     data_err = pd.to_numeric(df_err[p.n_value], errors='raise')
            #     testres = a.chauvenet(data_err)
            #     df_orig.loc[testres[testres==True].index, 'chauvenet'] = 'F'
            #
            # # выставим значения ошибки структуры
            # df_orig['data_structure'] = ''
            # df_orig.loc[error_idx['nan'], 'data_structure'] = 'НЕ ЧИСЛО'
            # df_orig.loc[error_idx['no_nucleos'], 'data_structure'] += 'нет ядра'
            # df_orig.loc[error_idx['no_nucleori'], 'data_structure'] += 'нет ядрышек'
            # df_orig.loc[error_idx['additional_nucleos'], 'data_structure'] += 'лишнее ядро'
            #
            # #====
            # # сохранить подсвеченые данные
            # outerr_filepath = os.path.relpath(dir)
            # outerr_filepath = os.path.join(p.p_data_highlights, outerr_filepath)
            # os.makedirs(outerr_filepath, exist_ok=True)
            # outerr_filename = os.path.splitext(in_file)[0] + '.xls'
            # outerr_filename = os.path.join(outerr_filepath, outerr_filename)
            # df_orig.index = df_orig.index+2
            # df_orig.to_excel(outerr_filename, sheet_name=p.sheet_name)
            # #====
            df[p.n_value] = df[p.n_value].astype('float64')
            data.append(df)
        except ValueError as ex:
            print("ВНИМАНИЕ! Ошибка при обработке файла! "
                  "В значениях содержатся нечисловые значения!")
            print("file name = " + in_filepath)
            quit(1)
        except Exception as ex:
            print("ВНИМАНИЕ! Ошибка при обработке файла!")
            print("file name = " + in_filepath)
            ##print(ex)
            raise
            quit(1)

    
    if len(data) > 0:
        return pd.concat(data)
    else:
        return None


def compare_dirs_structure(dir1, dir2):
    dirs1 = os.listdir(dir1)
    dirs1 = list(filter(lambda x: os.path.isdir(os.path.join(dir1, x)), dirs1))
    dirs2 = os.listdir(dir2)
    dirs2 = list(filter(lambda x: os.path.isdir(os.path.join(dir2, x)), dirs2))

    rc = []
    rcd = []

    for d in dirs1:
        if d not in dirs2:
            rc.append((d, dir2))
        else:

            check_dir1 = os.path.join(dir1, d)
            check_dir2 = os.path.join(dir2, d)
            rcd.extend(compare_dirs_structure(check_dir1, check_dir2))

    rc.extend(rcd)
    return rc
